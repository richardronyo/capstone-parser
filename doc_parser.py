"""
Document parser for extracting chapters from various document formats.
Supports: PDF, DOCX, Markdown, reStructuredText, and plain text.
"""

import re
from typing import List, Dict, Any
from pathlib import Path

# PDF processing
from pypdf import PdfReader
import pdfplumber

# DOCX processing
from docx import Document as DocxDocument

# Markdown and RST
import markdown
from docutils.core import publish_doctree
from docutils.io import StringInput


# ---------------------------------------------------------------------------
# Chapter detection patterns
# ---------------------------------------------------------------------------

# Matches common chapter patterns:
# - "Chapter 1", "Chapter One", "CHAPTER 1:"
# - "1. Introduction", "1 Introduction"
# - "Part I", "Part One"
# - Numbered sections like "1.0 Overview"
CHAPTER_PATTERNS = [
    r'^Chapter\s+\d+',                    # Chapter 1, Chapter 23
    r'^Chapter\s+[IVX]+',                 # Chapter I, Chapter XIV
    r'^Chapter\s+\w+',                    # Chapter One, Chapter Introduction
    r'^CHAPTER\s+\d+',                    # CHAPTER 1
    r'^Part\s+\d+',                       # Part 1
    r'^Part\s+[IVX]+',                    # Part I
    r'^Part\s+\w+',                       # Part One
    r'^\d+\.\s+[A-Z]',                    # 1. Introduction (capitalized)
    r'^\d+\s+[A-Z][a-z]+',                # 1 Introduction
    r'^\d+\.\d+\s+[A-Z]',                 # 1.0 Introduction
]

# Markdown headers (# Chapter, ## Section)
MD_HEADER_PATTERN = r'^#{1,2}\s+(.+)$'

# Heading detection for plain text - lines that are:
# - Short (< 60 chars)
# - Followed by blank line or underline
# - Start with capital or number
HEADING_HEURISTIC = r'^[A-Z0-9#][\w\s\-:]{2,60}$'


def is_chapter_heading(line: str) -> bool:
    """Check if a line looks like a chapter heading."""
    line = line.strip()
    if not line:
        return False
    
    # Try explicit patterns first
    for pattern in CHAPTER_PATTERNS:
        if re.match(pattern, line, re.IGNORECASE):
            return True
    
    # Fallback to heuristic for plain text
    if re.match(HEADING_HEURISTIC, line):
        return True
    
    return False


# ---------------------------------------------------------------------------
# PDF parsing
# ---------------------------------------------------------------------------

def parse_pdf(file_path: str) -> Dict[str, Any]:
    """
    Extract chapters from a PDF using text extraction and heuristics.
    
    Strategy:
    1. Use pdfplumber for better text extraction
    2. Detect chapter headings by font size (if available) or pattern matching
    3. Group pages into chapters
    """
    chapters = []
    current_chapter = None
    
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            lines = text.split('\n')
            
            for line in lines:
                line_stripped = line.strip()
                
                # Check if this line is a chapter heading
                if is_chapter_heading(line_stripped):
                    # Save previous chapter
                    if current_chapter:
                        chapters.append(current_chapter)
                    
                    # Start new chapter
                    current_chapter = {
                        "title": line_stripped,
                        "start_page": page_num,
                        "content": "",
                    }
                elif current_chapter:
                    current_chapter["content"] += line + "\n"
            
            # Add page separator for multi-page chapters
            if current_chapter:
                current_chapter["content"] += f"\n[Page {page_num}]\n\n"
        
        # Save final chapter
        if current_chapter:
            chapters.append(current_chapter)
    
    # If no chapters found, treat entire document as one chapter
    if not chapters:
        with pdfplumber.open(file_path) as pdf:
            full_text = ""
            for page in pdf.pages:
                full_text += (page.extract_text() or "") + "\n\n"
            
            chapters = [{
                "title": "Document",
                "start_page": 1,
                "content": full_text.strip(),
            }]
    
    return {
        "format": "pdf",
        "chapters": chapters,
        "summary": {
            "total_chapters": len(chapters),
            "chapter_titles": [c["title"] for c in chapters],
        }
    }


# ---------------------------------------------------------------------------
# DOCX parsing
# ---------------------------------------------------------------------------

def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    Extract chapters from a DOCX file using paragraph styles.
    
    Strategy:
    1. Look for Heading 1 and Heading 2 styles as chapter markers
    2. Fallback to pattern matching if no styles found
    """
    doc = DocxDocument(file_path)
    chapters = []
    current_chapter = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Check if this is a heading (Heading1, Heading2, or Title)
        is_heading = (
            para.style.name in ('Heading 1', 'Heading 2', 'Title') or
            is_chapter_heading(text)
        )
        
        if is_heading and text:
            # Save previous chapter
            if current_chapter:
                chapters.append(current_chapter)
            
            # Start new chapter
            current_chapter = {
                "title": text,
                "style": para.style.name,
                "content": "",
            }
        elif current_chapter and text:
            current_chapter["content"] += text + "\n\n"
    
    # Save final chapter
    if current_chapter:
        chapters.append(current_chapter)
    
    # If no chapters found, treat entire document as one chapter
    if not chapters:
        full_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        chapters = [{
            "title": "Document",
            "style": "Normal",
            "content": full_text,
        }]
    
    return {
        "format": "docx",
        "chapters": chapters,
        "summary": {
            "total_chapters": len(chapters),
            "chapter_titles": [c["title"] for c in chapters],
        }
    }


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------

def parse_markdown(file_path: str) -> Dict[str, Any]:
    """
    Extract chapters from a Markdown file using headers.
    
    Strategy:
    1. Parse headers (# and ##)
    2. Treat each top-level header as a chapter
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    chapters = []
    current_chapter = None
    
    for line in lines:
        # Check for markdown headers
        header_match = re.match(r'^(#{1,2})\s+(.+)$', line)
        
        if header_match:
            level = len(header_match.group(1))
            title = header_match.group(2).strip()
            
            # Only treat # and ## as chapter boundaries
            if level <= 2:
                # Save previous chapter
                if current_chapter:
                    chapters.append(current_chapter)
                
                # Start new chapter
                current_chapter = {
                    "title": title,
                    "level": level,
                    "content": "",
                }
            else:
                # Include lower-level headers in content
                if current_chapter:
                    current_chapter["content"] += line + "\n"
        elif current_chapter:
            current_chapter["content"] += line + "\n"
    
    # Save final chapter
    if current_chapter:
        chapters.append(current_chapter)
    
    # If no chapters found, treat entire document as one chapter
    if not chapters:
        chapters = [{
            "title": "Document",
            "level": 1,
            "content": content.strip(),
        }]
    
    return {
        "format": "markdown",
        "chapters": chapters,
        "summary": {
            "total_chapters": len(chapters),
            "chapter_titles": [c["title"] for c in chapters],
        }
    }


# ---------------------------------------------------------------------------
# reStructuredText parsing
# ---------------------------------------------------------------------------

def parse_rst(file_path: str) -> Dict[str, Any]:
    """
    Extract chapters from a reStructuredText file.
    
    Strategy:
    1. Parse RST using docutils
    2. Extract section nodes as chapters
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse RST into a document tree
    doctree = publish_doctree(
        content,
        source_class=StringInput,
        settings_overrides={'report_level': 5}  # Suppress warnings
    )
    
    chapters = []
    
    # Walk through sections in the document tree
    for section in doctree.traverse(condition=lambda n: n.tagname == 'section'):
        # Get section title
        title_node = section.next_node(condition=lambda n: n.tagname == 'title')
        title = title_node.astext() if title_node else "Untitled"
        
        # Get section content (all text nodes)
        content_parts = []
        for node in section.traverse(condition=lambda n: hasattr(n, 'astext')):
            if node.tagname not in ('title', 'section'):
                text = node.astext().strip()
                if text:
                    content_parts.append(text)
        
        chapters.append({
            "title": title,
            "content": "\n\n".join(content_parts),
        })
    
    # If no sections found, treat entire document as one chapter
    if not chapters:
        chapters = [{
            "title": "Document",
            "content": doctree.astext(),
        }]
    
    return {
        "format": "rst",
        "chapters": chapters,
        "summary": {
            "total_chapters": len(chapters),
            "chapter_titles": [c["title"] for c in chapters],
        }
    }


# ---------------------------------------------------------------------------
# Plain text parsing
# ---------------------------------------------------------------------------

def parse_txt(file_path: str) -> Dict[str, Any]:
    """
    Extract chapters from a plain text file using heuristics.
    
    Strategy:
    1. Look for lines that match chapter patterns
    2. Look for lines followed by blank lines (potential headings)
    3. Use line length and capitalization heuristics
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    chapters = []
    current_chapter = None
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        # Check if this looks like a heading
        if is_chapter_heading(line_stripped):
            # Confirm it's followed by blank line or content (not just noise)
            is_heading = False
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # Heading if followed by blank line or substantial content
                if not next_line or len(next_line) > 20:
                    is_heading = True
            
            if is_heading:
                # Save previous chapter
                if current_chapter:
                    chapters.append(current_chapter)
                
                # Start new chapter
                current_chapter = {
                    "title": line_stripped,
                    "content": "",
                }
                continue
        
        # Add content to current chapter
        if current_chapter:
            current_chapter["content"] += line
    
    # Save final chapter
    if current_chapter:
        chapters.append(current_chapter)
    
    # If no chapters found, treat entire document as one chapter
    if not chapters:
        with open(file_path, 'r', encoding='utf-8') as f:
            full_text = f.read()
        
        chapters = [{
            "title": Path(file_path).stem,
            "content": full_text,
        }]
    
    return {
        "format": "txt",
        "chapters": chapters,
        "summary": {
            "total_chapters": len(chapters),
            "chapter_titles": [c["title"] for c in chapters],
        }
    }


# ---------------------------------------------------------------------------
# Main dispatcher
# ---------------------------------------------------------------------------

def parse_document(file_path: str) -> Dict[str, Any]:
    """
    Parse a document and extract chapters.
    
    Supported formats: PDF, DOCX, MD, RST, TXT
    
    Returns:
        {
            "format": "pdf|docx|markdown|rst|txt",
            "chapters": [
                {
                    "title": "Chapter title",
                    "content": "Chapter content...",
                    "start_page": 1,  # PDF only
                    "level": 1,       # Markdown only
                    "style": "Heading 1"  # DOCX only
                },
                ...
            ],
            "summary": {
                "total_chapters": 5,
                "chapter_titles": ["Chapter 1", "Chapter 2", ...]
            }
        }
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext in (".md", ".markdown"):
        return parse_markdown(file_path)
    elif ext == ".rst":
        return parse_rst(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported format: {ext}. Supported: .pdf, .docx, .md, .rst, .txt")
